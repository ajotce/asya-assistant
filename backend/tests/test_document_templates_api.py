from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import Any

from docx import Document

from app.main import app
from app.storage.runtime import blob_storage
from tests.auth_helpers import build_authed_client


@dataclass
class _FakeResponse:
    content: bytes = b"%PDF-1.7"

    def raise_for_status(self) -> None:
        return


def _make_docx_template() -> bytes:
    doc = Document()
    doc.add_paragraph("VIN: {{vin}}")
    doc.add_paragraph("Клиент: {{client_name}}")
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def test_document_templates_crud_and_fill_both_files(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("DOCUMENTS_CONVERTER_ENABLED", "true")
    monkeypatch.setenv("DOCUMENTS_CONVERTER_URL", "http://libreoffice:3000")

    def _post(url: str, *args: Any, **kwargs: Any) -> _FakeResponse:
        assert url.endswith("/convert")
        return _FakeResponse()

    monkeypatch.setattr("app.services.docx_to_pdf_converter.httpx.post", _post)

    client = build_authed_client(tmp_path, monkeypatch, "templates@example.com", "Templates")

    template_key = "templates/geely-guarantee.docx"
    blob_storage.put_bytes(template_key, _make_docx_template())

    created = client.post(
        "/api/document-templates",
        json={
            "name": "Гарантия Geely",
            "description": "Шаблон гарантийного документа",
            "provider": "google_drive",
            "file_id": template_key,
            "fields": [
                {
                    "key": "vin",
                    "label": "VIN",
                    "type": "vin",
                    "required": True,
                    "validation": None,
                },
                {
                    "key": "client_name",
                    "label": "ФИО клиента",
                    "type": "text",
                    "required": True,
                    "validation": None,
                },
            ],
            "output_settings": {
                "format": "both",
                "filename": "geely-warranty",
            },
        },
    )
    assert created.status_code == 201
    template_id = created.json()["id"]

    listed = client.get("/api/document-templates")
    assert listed.status_code == 200
    assert any(item["id"] == template_id for item in listed.json())

    filled = client.post(
        f"/api/document-templates/{template_id}/fill",
        json={
            "values": {
                "vin": "1HGCM82633A004352",
                "client_name": "Иван Иванов",
            }
        },
    )
    assert filled.status_code == 200
    files = filled.json()["files"]
    assert len(files) == 2
    assert {f["content_type"] for f in files} == {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/pdf",
    }

    patched = client.patch(
        f"/api/document-templates/{template_id}",
        json={
            "description": "Обновлённый шаблон",
            "output_settings": {"format": "docx", "filename": "geely-docx-only"},
        },
    )
    assert patched.status_code == 200
    assert patched.json()["description"] == "Обновлённый шаблон"

    deleted = client.delete(f"/api/document-templates/{template_id}")
    assert deleted.status_code == 204

    app.dependency_overrides.clear()

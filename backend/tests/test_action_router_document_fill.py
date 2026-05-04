from __future__ import annotations

from io import BytesIO

from docx import Document
from sqlalchemy.orm import Session

from app.db.models.common import UserRole, UserStatus
from app.db.models.user import User
from app.repositories.document_template_repository import DocumentTemplateRepository
from app.services.action_router import ActionRouter
from app.services.file_storage_service import FileStorageService
from tests.auth_helpers import setup_test_db


def _docx_template_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("Клиент {{client_name}}")
    doc.add_paragraph("VIN {{vin}}")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_action_router_document_fill_flow(tmp_path, monkeypatch) -> None:
    _, engine = setup_test_db(tmp_path, monkeypatch)
    with Session(bind=engine) as session:
        user = User(
            email="doc-router@example.com",
            display_name="Doc Router",
            password_hash="hash",
            role=UserRole.USER,
            status=UserStatus.ACTIVE,
        )
        session.add(user)
        session.flush()

        repo = DocumentTemplateRepository(session)
        repo.create(
            user_id=user.id,
            name="Гарантия Geely",
            description="",
            provider="google_drive",
            file_id="template-1",
            fields=[
                {"key": "client_name", "label": "ФИО", "type": "text", "required": True},
                {"key": "vin", "label": "VIN", "type": "vin", "required": True},
            ],
            output_settings={"filename": "warranty"},
        )
        session.flush()

        def _fake_read(self: FileStorageService, *, provider: str, item_id: str) -> bytes:
            return _docx_template_bytes()

        monkeypatch.setattr(FileStorageService, "read", _fake_read)

        pending_store: dict = {}
        router = ActionRouter(session, pending_store=pending_store)

        ask_missing = router.handle(
            user_id=user.id,
            session_id="chat-1",
            message="Заполни шаблон Гарантия Geely client_name=Иванов Иван",
        )
        assert ask_missing.handled is True
        assert "Missing fields" in ask_missing.message
        assert ask_missing.pending_action_id is not None

        done = router.handle(
            user_id=user.id,
            session_id="chat-1",
            message="vin=XW8ZZZ1BZGG123456",
        )
        assert done.handled is True
        assert "заполнен" in done.message
        assert "tmp" in done.message

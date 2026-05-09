from io import BytesIO

from docx import Document

from app.services.document_fill_service import DocumentFillService


def _build_template() -> bytes:
    doc = Document()
    doc.add_paragraph("Клиент: {{client_name}}")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "VIN"
    table.rows[0].cells[1].text = "{{vin}}"
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Паспорт: {{passport_number}}"
    section.footer.paragraphs[0].text = "Телефон: {{client_phone}}"

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def test_fill_template_replaces_placeholders_in_all_sections() -> None:
    service = DocumentFillService()
    filled = service.fill_template(
        _build_template(),
        {
            "client_name": "Иван Иванов",
            "vin": "XTA210990Y1234567",
            "passport_number": "1234 567890",
            "client_phone": "+79990001122",
        },
    )

    parsed = Document(BytesIO(filled))
    assert "Клиент: Иван Иванов" in [p.text for p in parsed.paragraphs]
    assert parsed.tables[0].rows[0].cells[1].text == "XTA210990Y1234567"
    assert parsed.sections[0].header.paragraphs[0].text == "Паспорт: 1234 567890"
    assert parsed.sections[0].footer.paragraphs[0].text == "Телефон: +79990001122"

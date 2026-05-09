from __future__ import annotations

from io import BytesIO

from docx import Document


class DocumentFillService:
    """Fill DOCX templates by replacing {{FieldKey}} placeholders."""

    def fill_template(self, template_bytes: bytes, values: dict[str, str]) -> bytes:
        doc = Document(BytesIO(template_bytes))
        replacements = {f"{{{{{key}}}}}": value for key, value in values.items()}

        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)

        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, replacements)

            for paragraph in section.footer.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, replacements)

        out = BytesIO()
        doc.save(out)
        return out.getvalue()

    @staticmethod
    def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
        # Replace at paragraph level to support placeholders split across runs.
        text = paragraph.text
        updated = text
        for placeholder, value in replacements.items():
            updated = updated.replace(placeholder, value)
        if updated == text:
            return

        if paragraph.runs:
            paragraph.runs[0].text = updated
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = updated

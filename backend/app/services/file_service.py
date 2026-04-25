from __future__ import annotations

from pathlib import Path
from typing import List
from uuid import uuid4

from fastapi import UploadFile

from app.core.config import Settings
from app.models.schemas import SessionUploadedFileInfo
from app.services.vsellm_client import VseLLMClient, VseLLMError
from app.storage.file_store import SessionFileStore, StoredSessionFile
from app.storage.session_store import SessionStore
from app.storage.usage_store import UsageStore
from app.storage.vector_store import SessionVectorStore, StoredChunkVector

_ALLOWED_DOC_EXTENSIONS = {".pdf", ".docx", ".xlsx"}
_ALLOWED_IMAGE_EXTENSIONS = {
    ".jpg",
    ".jpeg",
    ".png",
    ".gif",
    ".bmp",
    ".webp",
    ".tif",
    ".tiff",
    ".heic",
}
_ALLOWED_IMAGE_CONTENT_PREFIX = "image/"
_READ_CHUNK_SIZE = 1024 * 1024
_MAX_PREVIEW_CHARS = 200_000
_CHUNK_SIZE_CHARS = 1200
_CHUNK_OVERLAP_CHARS = 200


class FileValidationError(Exception):
    def __init__(self, user_message: str, status_code: int = 400) -> None:
        super().__init__(user_message)
        self.user_message = user_message
        self.status_code = status_code


class FileService:
    def __init__(
        self,
        settings: Settings,
        session_store: SessionStore,
        file_store: SessionFileStore,
        vector_store: SessionVectorStore,
        vsellm_client: VseLLMClient,
        usage_store: UsageStore | None = None,
    ) -> None:
        self._settings = settings
        self._session_store = session_store
        self._file_store = file_store
        self._vector_store = vector_store
        self._vsellm_client = vsellm_client
        self._usage_store = usage_store

    async def upload_files(self, session_id: str, files: List[UploadFile]) -> List[SessionUploadedFileInfo]:
        if not self._session_store.has_session(session_id):
            raise FileValidationError("Сессия не найдена.", status_code=404)
        if not files:
            raise FileValidationError("Не выбраны файлы для загрузки.")
        if len(files) > self._settings.max_files_per_message:
            raise FileValidationError(
                f"Можно загрузить не более {self._settings.max_files_per_message} файлов за одно сообщение."
            )

        pending_files: List[StoredSessionFile] = []
        try:
            for upload in files:
                pending = await self._save_single_upload(session_id=session_id, upload=upload)
                pending_files.append(pending)
                self._index_document_chunks(session_id=session_id, stored_file=pending)
        except (FileValidationError, VseLLMError) as exc:
            self._cleanup_pending_files(session_id=session_id, files=pending_files)
            if isinstance(exc, FileValidationError):
                raise
            raise FileValidationError(user_message=exc.user_message, status_code=exc.status_code) from exc
        finally:
            for upload in files:
                await upload.close()

        self._file_store.register_files(session_id, pending_files)
        for pending in pending_files:
            self._session_store.bind_file(session_id=session_id, file_id=pending.file_id)

        return [
            SessionUploadedFileInfo(
                file_id=file.file_id,
                filename=file.filename,
                content_type=file.content_type,
                size_bytes=file.size_bytes,
            )
            for file in pending_files
        ]

    async def _save_single_upload(self, session_id: str, upload: UploadFile) -> StoredSessionFile:
        filename = Path(upload.filename or "").name.strip()
        if not filename:
            raise FileValidationError("У файла отсутствует имя.")

        content_type = (upload.content_type or "application/octet-stream").strip()
        self._validate_supported_type(filename=filename, content_type=content_type)

        file_id = str(uuid4())
        session_dir = self._file_store.session_dir(session_id)
        session_dir.mkdir(parents=True, exist_ok=True)
        target_path = session_dir / f"{file_id}_{filename}"

        max_file_size_bytes = self._settings.max_file_size_mb * 1024 * 1024
        total_size = 0

        with target_path.open("wb") as out:
            while True:
                chunk = await upload.read(_READ_CHUNK_SIZE)
                if not chunk:
                    break
                total_size += len(chunk)
                if total_size > max_file_size_bytes:
                    out.close()
                    target_path.unlink(missing_ok=True)
                    raise FileValidationError(
                        f"Файл '{filename}' превышает лимит {self._settings.max_file_size_mb} МБ."
                    )
                out.write(chunk)

        if total_size == 0:
            target_path.unlink(missing_ok=True)
            raise FileValidationError(f"Файл '{filename}' пустой.")

        self._validate_image_payload_if_needed(
            filename=filename,
            content_type=content_type,
            target_path=target_path,
        )

        try:
            extracted_text = self._extract_text(filename=filename, target_path=target_path)
        except FileValidationError:
            target_path.unlink(missing_ok=True)
            raise

        return StoredSessionFile(
            file_id=file_id,
            session_id=session_id,
            filename=filename,
            content_type=content_type,
            size_bytes=total_size,
            path=str(target_path),
            extracted_text=extracted_text,
        )

    def _index_document_chunks(self, session_id: str, stored_file: StoredSessionFile) -> None:
        if not stored_file.extracted_text:
            self._vector_store.upsert_file_chunks(session_id=session_id, file_id=stored_file.file_id, chunks=[])
            return

        chunks_text = self._chunk_text(stored_file.extracted_text)
        if not chunks_text:
            raise FileValidationError(f"Файл '{stored_file.filename}' не содержит данных для индексирования.")

        vectors, embeddings_usage = self._get_embeddings_with_usage(chunks_text)
        if self._usage_store is not None and embeddings_usage is not None:
            self._usage_store.record_embeddings_usage(session_id=session_id, usage=embeddings_usage)
        indexed_chunks = [
            StoredChunkVector(
                chunk_id=f"{stored_file.file_id}:{index}",
                file_id=stored_file.file_id,
                filename=stored_file.filename,
                text=chunk_text,
                embedding=vector,
            )
            for index, (chunk_text, vector) in enumerate(zip(chunks_text, vectors), start=1)
        ]
        self._vector_store.upsert_file_chunks(
            session_id=session_id,
            file_id=stored_file.file_id,
            chunks=indexed_chunks,
        )

    def _get_embeddings_with_usage(self, texts: list[str]) -> tuple[list[list[float]], dict | None]:
        if hasattr(self._vsellm_client, "get_embeddings_with_usage"):
            result = self._vsellm_client.get_embeddings_with_usage(texts)
            return result.vectors, result.usage
        return self._vsellm_client.get_embeddings(texts), None

    def _cleanup_pending_files(self, session_id: str, files: List[StoredSessionFile]) -> None:
        for file in files:
            Path(file.path).unlink(missing_ok=True)
            self._vector_store.delete_file_chunks(session_id=session_id, file_id=file.file_id)

    @staticmethod
    def _validate_supported_type(filename: str, content_type: str) -> None:
        suffix = Path(filename).suffix.lower()
        if suffix in _ALLOWED_DOC_EXTENSIONS:
            return
        if suffix in _ALLOWED_IMAGE_EXTENSIONS:
            return
        if content_type.lower().startswith(_ALLOWED_IMAGE_CONTENT_PREFIX):
            return
        raise FileValidationError(
            f"Файл '{filename}' не поддерживается. Разрешены типы: PDF, DOCX, XLSX и изображения."
        )

    def _extract_text(self, filename: str, target_path: Path) -> str | None:
        suffix = target_path.suffix.lower()
        if suffix == ".pdf":
            return self._extract_pdf_text(filename=filename, target_path=target_path)
        if suffix == ".docx":
            return self._extract_docx_text(filename=filename, target_path=target_path)
        if suffix == ".xlsx":
            return self._extract_xlsx_text(filename=filename, target_path=target_path)
        return None

    def _validate_image_payload_if_needed(self, filename: str, content_type: str, target_path: Path) -> None:
        suffix = target_path.suffix.lower()
        if suffix not in _ALLOWED_IMAGE_EXTENSIONS and not content_type.lower().startswith(_ALLOWED_IMAGE_CONTENT_PREFIX):
            return
        try:
            from PIL import Image
        except ImportError as exc:  # pragma: no cover
            raise FileValidationError("Сервис временно не готов к валидации изображений.", status_code=503) from exc

        try:
            with Image.open(target_path) as img:
                img.verify()
        except Exception as exc:
            target_path.unlink(missing_ok=True)
            raise FileValidationError(f"Изображение '{filename}' повреждено или имеет некорректный формат.") from exc

    def _extract_pdf_text(self, filename: str, target_path: Path) -> str:
        try:
            import fitz  # PyMuPDF
        except ImportError as exc:  # pragma: no cover
            raise FileValidationError("Сервис временно не готов к обработке PDF-файлов.", status_code=503) from exc

        try:
            doc = fitz.open(target_path)
        except Exception as exc:
            raise FileValidationError(f"Не удалось прочитать PDF '{filename}'. Файл может быть повреждён.") from exc

        parts: list[str] = []
        try:
            for page_index, page in enumerate(doc, start=1):
                text = (page.get_text("text") or "").strip()
                if text:
                    parts.append(f"[Страница {page_index}]\n{text}")
        finally:
            doc.close()

        combined = "\n\n".join(parts).strip()
        if not combined:
            raise FileValidationError(f"PDF '{filename}' не содержит текстовых данных.")
        return combined[:_MAX_PREVIEW_CHARS]

    def _extract_docx_text(self, filename: str, target_path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover
            raise FileValidationError("Сервис временно не готов к обработке DOCX-файлов.", status_code=503) from exc

        try:
            doc = Document(str(target_path))
        except Exception as exc:
            raise FileValidationError(f"Не удалось прочитать DOCX '{filename}'. Файл может быть повреждён.") from exc

        lines: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)

        for table_idx, table in enumerate(doc.tables, start=1):
            lines.append(f"[Таблица {table_idx}]")
            for row_idx, row in enumerate(table.rows, start=1):
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rendered = " | ".join(cell if cell else "—" for cell in cells)
                    lines.append(f"Строка {row_idx}: {rendered}")

        combined = "\n".join(lines).strip()
        if not combined:
            raise FileValidationError(f"DOCX '{filename}' не содержит текстовых данных.")
        return combined[:_MAX_PREVIEW_CHARS]

    def _extract_xlsx_text(self, filename: str, target_path: Path) -> str:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:  # pragma: no cover
            raise FileValidationError("Сервис временно не готов к обработке XLSX-файлов.", status_code=503) from exc

        try:
            workbook = load_workbook(filename=str(target_path), read_only=True, data_only=True)
        except Exception as exc:
            raise FileValidationError(f"Не удалось прочитать XLSX '{filename}'. Файл может быть повреждён.") from exc

        try:
            lines: list[str] = []
            has_any_data = False
            for sheet in workbook.worksheets:
                lines.append(f"[Лист: {sheet.title}]")
                has_data = False
                for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                    if not row:
                        continue
                    normalized = [self._normalize_cell_value(value) for value in row]
                    if not any(normalized):
                        continue
                    has_data = True
                    cells = [f"{idx + 1}={value}" for idx, value in enumerate(normalized) if value]
                    if cells:
                        lines.append(f"Строка {row_index}: " + " | ".join(cells))
                        has_any_data = True
                if not has_data:
                    lines.append("(пустой лист)")
            combined = "\n".join(lines).strip()
        finally:
            workbook.close()

        if not combined or not has_any_data:
            raise FileValidationError(f"XLSX '{filename}' не содержит данных для извлечения.")
        return combined[:_MAX_PREVIEW_CHARS]

    @staticmethod
    def _normalize_cell_value(value: object) -> str:
        if value is None:
            return ""
        if isinstance(value, float):
            if value.is_integer():
                return str(int(value))
            return str(value)
        return str(value).strip()

    @staticmethod
    def _chunk_text(text: str, chunk_size: int = _CHUNK_SIZE_CHARS, overlap: int = _CHUNK_OVERLAP_CHARS) -> list[str]:
        normalized = " ".join(text.split())
        if not normalized:
            return []
        if len(normalized) <= chunk_size:
            return [normalized]

        chunks: list[str] = []
        start = 0
        while start < len(normalized):
            end = min(start + chunk_size, len(normalized))
            chunk = normalized[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end >= len(normalized):
                break
            start = max(0, end - overlap)
        return chunks
